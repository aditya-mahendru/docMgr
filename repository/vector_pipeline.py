"""
Vector Pipeline for Document Ingestion
Handles document chunking, embedding generation, and vector storage
"""

from typing import List, Dict, Any, Optional
import markdown
import tiktoken
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
from dotenv import load_dotenv

load_dotenv()

class VectorPipeline:
    def __init__(self):
        """Initialize the vector pipeline with sentence-transformers and ChromaDB"""
        # Initialize sentence-transformers model
        print("🔄 Loading sentence transformer model...")
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        print("✅ Sentence transformer model loaded successfully")
        
        # Initialize ChromaDB
        self.chroma_client = chromadb.PersistentClient(
            path="./chroma_db",
            settings=Settings(anonymized_telemetry=False)
        )
        
        # Get or create collection
        self.collection = self.chroma_client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=self._count_tokens,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Tokenizer for counting tokens
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken"""
        return len(self.tokenizer.encode(text))
    
    def _extract_text_from_file(self, file_path: str, content_type: str) -> str:
        """Extract text content from different file types"""
        try:
            if content_type == "text/plain" or file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif content_type == "text/markdown" or file_path.endswith('.md'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                    # Convert markdown to plain text for processing
                    html = markdown.markdown(md_content)
                    # Simple HTML to text conversion (remove tags)
                    import re
                    text = re.sub(r'<[^>]+>', '', html)
                    return text
            
            elif content_type == "application/pdf" or file_path.endswith('.pdf'):
                return self._extract_text_from_pdf(file_path)
            
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_path.endswith('.docx'):
                return self._extract_text_from_docx(file_path)
            
            elif content_type.startswith("image/") or file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                return self._extract_text_from_image(file_path)
            
            else:
                raise ValueError(f"Unsupported file type: {content_type}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods for better results"""
        try:
            # Try pdfplumber first (better for complex layouts)
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text.strip())
                    if text_parts:
                        return '\n\n'.join(text_parts)
            except ImportError:
                pass
            
            # Fallback to PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text.strip())
                    if text_parts:
                        return '\n\n'.join(text_parts)
            except ImportError:
                pass
            
            # If both methods fail, raise error
            raise Exception("PDF text extraction failed - neither pdfplumber nor PyPDF2 available")
            
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # Extract text from headers and footers
            for section in doc.sections:
                for header in section.header.paragraphs:
                    if header.text.strip():
                        text_parts.append(header.text.strip())
                for footer in section.footer.paragraphs:
                    if footer.text.strip():
                        text_parts.append(footer.text.strip())
            
            if text_parts:
                return '\n\n'.join(text_parts)
            else:
                raise Exception("No text content found in DOCX file")
                
        except ImportError:
            raise Exception("python-docx library not available for DOCX processing")
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR and generate description using Groq API"""
        try:
            import os
            from PIL import Image
            import pytesseract
            import cv2
            from groq import Groq
            
            # Check if Groq API key is available
            groq_api_key = os.getenv("GROQ_API_KEY")
            if not groq_api_key:
                raise Exception("GROQ_API_KEY environment variable not set")
            
            # Initialize Groq client
            groq_client = Groq(api_key=groq_api_key)
            
            # Load and preprocess image
            image = cv2.imread(file_path)
            if image is None:
                raise Exception("Could not load image file")
            
            # Convert to grayscale for better OCR
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply preprocessing for better OCR results
            # Denoise
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Increase contrast
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(denoised)
            
            # OCR extraction
            ocr_text = pytesseract.image_to_string(enhanced, config='--psm 6')
            
            if not ocr_text.strip():
                # Try with different PSM modes if no text found
                ocr_text = pytesseract.image_to_string(enhanced, config='--psm 3')
            
            # Generate detailed description using Groq API
            prompt = f"""
            Analyze this OCR text from an image and provide a detailed, structured description.
            The image likely contains text, tables, bills, receipts, or other documents.
            
            OCR Text: {ocr_text}  # Limit to avoid token limits
            
            Please provide:
            1. A detailed description of what this document appears to be including items, amounts, dates, type of payment, etc.
            2. Key information extracted (dates, amounts, names, etc.)
            3. Any tables or structured data identified
            4. Overall document type and purpose
            
            Format your response as clear, searchable text that can be used for document retrieval.
            """
            
            try:
                response = groq_client.chat.completions.create(
                    model="openai/gpt-oss-20b",  # Using GPT OSS 20B for better performance
                    messages=[
                        {"role": "system", "content": "You are a document analysis expert. Provide clear, structured descriptions of documents based on OCR text."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=1000,
                    temperature=0.3
                )
                
                groq_description = response.choices[0].message.content.strip()
                
                # Combine OCR text and Groq description
                combined_text = f"OCR Text:\n{ocr_text}\n\nAI Description:\n{groq_description}"
                
                return combined_text
                
            except Exception as groq_error:
                print(f"Warning: Groq API call failed: {groq_error}")
                # Fallback to just OCR text if Groq fails
                return f"OCR Text:\n{ocr_text}\n\nNote: AI description generation failed"
                
        except ImportError as e:
            raise Exception(f"Required library not available: {str(e)}")
        except Exception as e:
            raise Exception(f"Error processing image: {str(e)}")
    
    def _chunk_text(self, text: str) -> List[str]:
        """Break text into chunks using LangChain splitter"""
        return self.text_splitter.split_text(text)
    
    def _generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks using sentence-transformers"""
        try:
            # Generate embeddings locally using sentence-transformers
            embeddings = self.embedding_model.encode(chunks, convert_to_tensor=False)
            # Convert to list of lists of floats
            if hasattr(embeddings, 'tolist'):
                return embeddings.tolist()
            elif isinstance(embeddings, list):
                return embeddings
            else:
                # Handle numpy array case
                import numpy as np
                if isinstance(embeddings, np.ndarray):
                    return embeddings.tolist()
                else:
                    # Fallback: convert to list
                    return [list(emb) for emb in embeddings]
        except Exception as e:
            raise Exception(f"Error generating embeddings: {str(e)}")
    
    def process_document(self, file_path: str, content_type: str, document_id: int, 
                        original_filename: str, description: Optional[str] = None) -> Dict[str, Any]:
        """Process a document through the vector pipeline"""
        try:
            # Extract text from file
            text = self._extract_text_from_file(file_path, content_type)
            
            # Chunk the text
            chunks = self._chunk_text(text)
            
            # Generate embeddings
            embeddings = self._generate_embeddings(chunks)
            
            # Store in vector database
            chunk_ids = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_id = f"{document_id}_{i}"
                chunk_ids.append(chunk_id)
                
                # Prepare metadata, ensuring no None values
                metadata = {
                    "document_id": document_id,
                    "chunk_index": i,
                    "original_filename": original_filename,
                    "content_type": content_type,
                    "description": description or "",  # Convert None to empty string
                    "chunk_size": self._count_tokens(chunk),
                    "total_chunks": len(chunks)
                }
                
                # Store chunk metadata and embedding
                self.collection.add(
                    embeddings=[embedding],
                    documents=[chunk],
                    metadatas=[metadata],
                    ids=[chunk_id]
                )
            
            return {
                "document_id": document_id,
                "original_filename": original_filename,
                "total_chunks": len(chunks),
                "total_tokens": sum(self._count_tokens(chunk) for chunk in chunks),
                "chunk_ids": chunk_ids,
                "status": "processed"
            }
            
        except Exception as e:
            return {
                "document_id": document_id,
                "original_filename": original_filename,
                "status": "error",
                "error": str(e)
            }
    
    def search_documents(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        """Search documents by semantic similarity"""
        try:
            # Generate embedding for query
            query_embedding = self._generate_embeddings([query])[0]
            
            # Search in vector database
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                include=["documents", "metadatas", "distances"]
            )
            
            # Format results
            formatted_results = []
            if results and 'ids' in results and results['ids'] and len(results['ids']) > 0:
                for i in range(len(results['ids'][0])):
                    formatted_results.append({
                        "chunk_id": results['ids'][0][i],
                        "content": results['documents'][0][i],
                        "metadata": results['metadatas'][0][i],
                        "similarity_score": 1 - results['distances'][0][i]  # Convert distance to similarity
                    })
            
            return formatted_results
            
        except Exception as e:
            raise Exception(f"Error searching documents: {str(e)}")
    
    def get_document_chunks(self, document_id: int) -> List[Dict[str, Any]]:
        """Get all chunks for a specific document"""
        try:
            results = self.collection.get(
                where={"document_id": document_id},
                include=["documents", "metadatas"]
            )
            
            chunks = []
            if results and 'ids' in results and results['ids']:
                for i in range(len(results['ids'])):
                    chunks.append({
                        "chunk_id": results['ids'][i],
                        "content": results['documents'][i],
                        "metadata": results['metadatas'][i]
                    })
            
            return sorted(chunks, key=lambda x: x['metadata']['chunk_index'])
            
        except Exception as e:
            raise Exception(f"Error retrieving document chunks: {str(e)}")
    
    def delete_document_chunks(self, document_id: int) -> bool:
        """Delete all chunks for a specific document"""
        try:
            # Get chunk IDs for the document
            results = self.collection.get(
                where={"document_id": document_id},
                include=[]
            )
            
            if results and 'ids' in results and results['ids']:
                self.collection.delete(ids=results['ids'])
            
            return True
            
        except Exception as e:
            raise Exception(f"Error deleting document chunks: {str(e)}")
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector collection"""
        try:
            count = self.collection.count()
            
            # Get sample metadata to understand structure
            sample = self.collection.get(limit=1, include=["metadatas"])
            
            return {
                "total_chunks": count,
                "collection_name": "documents",
                "sample_metadata": sample and 'metadatas' in sample and sample['metadatas'][0] if sample['metadatas'] else None
            }
            
        except Exception as e:
            return {"error": str(e)}
