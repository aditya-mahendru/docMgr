#!/usr/bin/env python3
"""
Test script for Vector Pipeline functionality
Tests document processing, search, and chunk retrieval
"""

import requests
import os
import json
from pathlib import Path

# API base URL
BASE_URL = "http://localhost:8000"

def create_test_documents():
    """Create test text, markdown, and PDF documents"""
    test_docs = []
    
    # Create test directory if it doesn't exist
    test_dir = Path("test_docs")
    test_dir.mkdir(exist_ok=True)
    
    # Test text document
    text_content = """
    This is a test document about artificial intelligence and machine learning.
    
    Artificial Intelligence (AI) is a branch of computer science that aims to create intelligent machines.
    These machines can perform tasks that typically require human intelligence, such as visual perception,
    speech recognition, decision-making, and language translation.
    
    Machine Learning is a subset of AI that focuses on algorithms and statistical models.
    It enables computers to improve their performance on a specific task through experience.
    """
    
    text_file = test_dir / "ai_ml_document.txt"
    with open(text_file, "w", encoding="utf-8") as f:
        f.write(text_content)
    test_docs.append(str(text_file))
    
    # Test markdown document
    md_content = """
    # Data Science Fundamentals
    
    ## Introduction
    Data science is an interdisciplinary field that uses scientific methods, processes, algorithms, and systems
    to extract knowledge and insights from structured and unstructured data.
    
    ## Key Components
    1. **Data Collection**: Gathering data from various sources
    2. **Data Cleaning**: Preparing and validating data
    3. **Data Analysis**: Exploring and understanding patterns
    4. **Data Visualization**: Presenting findings visually
    5. **Machine Learning**: Building predictive models
    
    ## Applications
    - Business Intelligence
    - Healthcare Analytics
    - Financial Modeling
    - Scientific Research
    """
    
    md_file = test_dir / "data_science.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(md_content)
    test_docs.append(str(md_file))
    
    # Test DOCX document
    try:
        from docx import Document
        
        docx_file = test_dir / "business_report.docx"
        doc = Document()
        
        # Add title
        doc.add_heading('Business Report: Q4 2024', 0)
        
        # Add paragraph
        doc.add_paragraph('This quarterly report summarizes our business performance and strategic initiatives.')
        
        # Add section
        doc.add_heading('Financial Performance', level=1)
        doc.add_paragraph('Revenue increased by 15% compared to Q3, driven by strong product adoption.')
        
        # Add table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Q3 2024'
        hdr_cells[2].text = 'Q4 2024'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Revenue'
        row_cells[1].text = '$2.1M'
        row_cells[2].text = '$2.4M'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Customers'
        row_cells[1].text = '1,250'
        row_cells[2].text = '1,450'
        
        doc.save(str(docx_file))
        test_docs.append(str(docx_file))
        print("✅ Created test DOCX document")
        
    except ImportError:
        print("⚠️  python-docx not available - skipping DOCX test")
        # Create a simple text file as DOCX substitute
        docx_substitute = test_dir / "business_report.txt"
        docx_content = """
        Business Report: Q4 2024
        
        This quarterly report summarizes our business performance and strategic initiatives.
        
        Financial Performance
        Revenue increased by 15% compared to Q3, driven by strong product adoption.
        
        Key Metrics:
        - Revenue: $2.4M (up from $2.1M)
        - Customers: 1,450 (up from 1,250)
        """
        with open(docx_substitute, "w", encoding="utf-8") as f:
            f.write(docx_content)
        test_docs.append(str(docx_substitute))
    
    # Test PDF document (create a simple text-based PDF)
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_file = test_dir / "research_paper.pdf"
        c = canvas.Canvas(str(pdf_file), pagesize=letter)
        
        # Add content to PDF
        c.drawString(100, 750, "Research Paper: Neural Networks")
        c.drawString(100, 720, "Abstract:")
        c.drawString(100, 700, "This paper explores the application of neural networks")
        c.drawString(100, 680, "in modern machine learning systems. We present")
        c.drawString(100, 660, "novel approaches to deep learning architecture")
        c.drawString(100, 640, "and demonstrate improved performance on benchmark")
        c.drawString(100, 620, "datasets.")
        
        c.drawString(100, 580, "1. Introduction")
        c.drawString(100, 560, "Neural networks have revolutionized the field of")
        c.drawString(100, 540, "artificial intelligence. Their ability to learn")
        c.drawString(100, 520, "complex patterns from data makes them suitable")
        c.drawString(100, 500, "for a wide range of applications.")
        
        c.save()
        test_docs.append(str(pdf_file))
        print("✅ Created test PDF document")
        
    except ImportError:
        print("⚠️  reportlab not available - skipping PDF test")
        # Create a simple text file as PDF substitute
        pdf_substitute = test_dir / "research_paper.txt"
        pdf_content = """
        Research Paper: Neural Networks
        
        Abstract:
        This paper explores the application of neural networks
        in modern machine learning systems. We present
        novel approaches to deep learning architecture
        and demonstrate improved performance on benchmark
        datasets.
        
        1. Introduction
        Neural networks have revolutionized the field of
        artificial intelligence. Their ability to learn
        complex patterns from data makes them suitable
        for a wide range of applications.
        """
        with open(pdf_substitute, "w", encoding="utf-8") as f:
            f.write(pdf_content)
        test_docs.append(str(pdf_substitute))
    
    return test_docs

def test_vector_pipeline():
    """Test the complete vector pipeline workflow"""
    print("🧪 Testing Vector Pipeline Functionality")
    print("=" * 60)
    
    # Check if vector pipeline is available
    try:
        response = requests.get(f"{BASE_URL}/")
        if response.status_code == 200:
            api_info = response.json()
            if not api_info.get("vector_pipeline", False):
                print("❌ Vector pipeline is not available. Check your dependencies installation.")
                return
            print("✅ Vector pipeline is available")
        else:
            print(f"❌ Failed to get API info: {response.status_code}")
            return
    except requests.exceptions.ConnectionError:
        print("❌ Could not connect to the API. Make sure the server is running on localhost:8000")
        return
    
    # Create test documents
    test_docs = create_test_documents()
    print(f"📝 Created {len(test_docs)} test documents")
    
    # Upload documents
    uploaded_docs = []
    for doc_path in test_docs:
        print(f"\n📤 Uploading {os.path.basename(doc_path)}...")
        
        with open(doc_path, "rb") as f:
            files = {"file": (os.path.basename(doc_path), f.read(), "text/plain")}
            data = {"description": f"Test document: {os.path.basename(doc_path)}"}
            
            response = requests.post(f"{BASE_URL}/api/documents/upload", files=files, data=data)
            
            if response.status_code == 200:
                doc_info = response.json()
                uploaded_docs.append(doc_info)
                print(f"✅ Uploaded successfully (ID: {doc_info['id']})")
                
                # Check if vector processing was successful
                if doc_info.get("processing_result"):
                    proc_result = doc_info["processing_result"]
                    if proc_result["status"] == "processed":
                        print(f"   📊 Vector processing: {proc_result['total_chunks']} chunks, {proc_result['total_tokens']} tokens")
                    else:
                        print(f"   ⚠️  Vector processing failed: {proc_result.get('error', 'Unknown error')}")
                else:
                    print("   ℹ️  No vector processing result available")
            else:
                print(f"❌ Upload failed: {response.status_code} - {response.text}")
    
    if not uploaded_docs:
        print("❌ No documents were uploaded successfully")
        return
    
    # Test vector statistics
    print(f"\n📊 Getting vector collection statistics...")
    try:
        response = requests.get(f"{BASE_URL}/api/vector/stats")
        if response.status_code == 200:
            stats = response.json()
            print(f"✅ Collection stats: {stats['total_chunks']} total chunks")
        else:
            print(f"❌ Failed to get stats: {response.status_code}")
    except Exception as e:
        print(f"❌ Error getting stats: {e}")
    
    # Test document chunks retrieval
    for doc in uploaded_docs:
        print(f"\n🔍 Getting chunks for document {doc['id']}...")
        try:
            response = requests.get(f"{BASE_URL}/api/documents/{doc['id']}/chunks")
            if response.status_code == 200:
                chunks = response.json()
                print(f"✅ Retrieved {len(chunks)} chunks")
                for i, chunk in enumerate(chunks[:2]):  # Show first 2 chunks
                    print(f"   Chunk {i+1}: {chunk['metadata']['chunk_size']} tokens")
                    print(f"   Content preview: {chunk['content'][:100]}...")
            else:
                print(f"❌ Failed to get chunks: {response.status_code}")
        except Exception as e:
            print(f"❌ Error getting chunks: {e}")
    
    # Test semantic search
    print(f"\n🔎 Testing semantic search...")
    search_queries = [
        "What is artificial intelligence?",
        "How does machine learning work?",
        "What are the key components of data science?",
        "Tell me about business intelligence applications"
    ]
    
    for query in search_queries:
        print(f"\n   Searching for: '{query}'")
        try:
            search_data = {"query": query, "n_results": 3}
            response = requests.post(f"{BASE_URL}/api/search", json=search_data)
            
            if response.status_code == 200:
                results = response.json()
                print(f"   ✅ Found {len(results)} results")
                for i, result in enumerate(results[:2]):  # Show top 2 results
                    score = result['similarity_score']
                    filename = result['metadata']['original_filename']
                    content_preview = result['content'][:80]
                    print(f"     {i+1}. {filename} (Score: {score:.3f})")
                    print(f"        {content_preview}...")
            else:
                print(f"   ❌ Search failed: {response.status_code}")
        except Exception as e:
            print(f"   ❌ Search error: {e}")
    
    # Clean up test documents
    print(f"\n🧹 Cleaning up test documents...")
    for doc in uploaded_docs:
        try:
            response = requests.delete(f"{BASE_URL}/api/documents/{doc['id']}")
            if response.status_code == 200:
                print(f"   ✅ Deleted document {doc['id']}")
            else:
                print(f"   ❌ Failed to delete document {doc['id']}")
        except Exception as e:
            print(f"   ❌ Error deleting document {doc['id']}: {e}")
    
    # Remove test files
    for doc_path in test_docs:
        try:
            os.remove(doc_path)
        except:
            pass
    
    # Remove test directory if empty
    try:
        test_dir = Path("test_docs")
        if test_dir.exists() and not any(test_dir.iterdir()):
            test_dir.rmdir()
    except:
        pass
    
    print(f"\n" + "=" * 60)
    print("🎉 Vector Pipeline Test Completed!")

if __name__ == "__main__":
    print("Document Manager API - Vector Pipeline Test")
    print("Using local sentence-transformers for embeddings")
    print("=" * 60)
    
    test_vector_pipeline()
