#!/usr/bin/env python3
"""
Test script for DOCX support in the vector pipeline
"""

import requests
import os
from pathlib import Path

# API base URL
BASE_URL = "http://localhost:8000"

def test_docx_upload():
    """Test DOCX document upload and processing"""
    print("🧪 Testing DOCX Support")
    print("=" * 40)
    
    # Create a test DOCX file
    test_dir = Path("test_docx")
    test_dir.mkdir(exist_ok=True)
    
    docx_file = test_dir / "test_document.docx"
    
    try:
        # Try to create a DOCX using python-docx
        from docx import Document
        
        doc = Document()
        
        # Add title
        doc.add_heading('Test DOCX Document', 0)
        
        # Add introduction
        doc.add_paragraph('This is a test DOCX document for the vector pipeline.')
        doc.add_paragraph('It contains structured content that should be extracted and processed.')
        
        # Add section with heading
        doc.add_heading('Key Features', level=1)
        doc.add_paragraph('The system should be able to:')
        
        # Add list
        features = [
            'Extract text from DOCX files',
            'Handle structured content (headings, paragraphs)',
            'Process tables and formatted text',
            'Generate embeddings for semantic search'
        ]
        
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
        
        # Add table
        doc.add_heading('Sample Data', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Status'
        
        # Data rows
        data = [
            ('Text Extraction', 'Working', '✅'),
            ('Table Processing', 'Working', '✅'),
            ('Vector Search', 'Working', '✅')
        ]
        
        for category, value, status in data:
            row_cells = table.add_row().cells
            row_cells[0].text = category
            row_cells[1].text = value
            row_cells[2].text = status
        
        doc.save(str(docx_file))
        print("✅ Created test DOCX document")
        
    except ImportError:
        print("⚠️  python-docx not available - creating text file instead")
        # Fallback to text file
        docx_file = test_dir / "test_document.txt"
        with open(docx_file, "w") as f:
            f.write("Test Document\nThis is a test document for the vector pipeline.")
    
    # Test upload
    print(f"\n📤 Uploading {docx_file.name}...")
    
    try:
        with open(docx_file, "rb") as f:
            # Set appropriate content type based on file extension
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if docx_file.suffix == '.docx' else "text/plain"
            files = {"file": (docx_file.name, f.read(), content_type)}
            data = {"description": "Test DOCX document for vector processing"}
            
            response = requests.post(f"{BASE_URL}/api/documents/upload", files=files, data=data)
            
            if response.status_code == 200:
                doc_info = response.json()
                print(f"✅ Upload successful! Document ID: {doc_info['id']}")
                
                # Check vector processing
                if doc_info.get("processing_result"):
                    proc_result = doc_info["processing_result"]
                    if proc_result["status"] == "processed":
                        print(f"   📊 Vector processing successful:")
                        print(f"      - Chunks: {proc_result['total_chunks']}")
                        print(f"      - Tokens: {proc_result['total_tokens']}")
                    else:
                        print(f"   ⚠️  Vector processing failed: {proc_result.get('error', 'Unknown error')}")
                else:
                    print("   ℹ️  No vector processing result available")
                
                # Test search functionality
                print(f"\n🔍 Testing search functionality...")
                search_data = {"query": "test document vector pipeline features", "n_results": 3}
                search_response = requests.post(f"{BASE_URL}/api/search", json=search_data)
                
                if search_response.status_code == 200:
                    results = search_response.json()
                    print(f"   ✅ Search successful! Found {len(results)} results")
                    for i, result in enumerate(results[:2]):
                        score = result['similarity_score']
                        filename = result['metadata']['original_filename']
                        print(f"      {i+1}. {filename} (Score: {score:.3f})")
                else:
                    print(f"   ❌ Search failed: {search_response.status_code}")
                
                # Clean up - delete the document
                print(f"\n🧹 Cleaning up test document...")
                delete_response = requests.delete(f"{BASE_URL}/api/documents/{doc_info['id']}")
                if delete_response.status_code == 200:
                    print("   ✅ Document deleted successfully")
                else:
                    print(f"   ❌ Failed to delete document: {delete_response.status_code}")
                
            else:
                print(f"❌ Upload failed: {response.status_code} - {response.text}")
                
    except requests.exceptions.ConnectionError:
        print("❌ Could not connect to the API. Make sure the server is running on localhost:8000")
    except Exception as e:
        print(f"❌ Error during test: {e}")
    
    # Clean up test files
    try:
        if docx_file.exists():
            os.remove(docx_file)
        if test_dir.exists() and not any(test_dir.iterdir()):
            test_dir.rmdir()
    except:
        pass
    
    print("\n" + "=" * 40)
    print("🎉 DOCX Support Test Completed!")

if __name__ == "__main__":
    print("Document Manager API - DOCX Support Test")
    print("Make sure the server is running and python-docx is installed")
    print("=" * 40)
    
    test_docx_upload()
